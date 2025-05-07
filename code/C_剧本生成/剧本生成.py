import pandas as pd
from docx import Document # python-docx
from docx.shared import Pt
from docx.oxml.ns import qn

# 读取CSV数据
input_path = "../B_小说解析/output/对话剧本_结构化数据_纯基础版.csv"
df = pd.read_csv(input_path)

# 创建Word文档
doc = Document()
doc.styles['Normal'].font.name = u'宋体'  # 设置中文字体
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 遍历每一行，添加到文档中
for _, row in df.iterrows():
    role = str(row["role"]).strip()
    dialogue = str(row["dialogue"]).strip()

    dialogue.replace(r'\n\n', '')

    # 跳过空dialogue
    if not dialogue or dialogue.startswith("(生成失败"):
        continue

    if role in ["", "旁白"]:
        para = doc.add_paragraph()
        run = para.add_run(dialogue)
        run.font.size = Pt(10)
        run.bold = False
    else:
        para = doc.add_paragraph()
        run = para.add_run(f"{role}：{dialogue}")
        run.font.size = Pt(12)
        run.bold = True

# 保存文档
output_path = "对比剧本/剧本输出_纯基础版.docx"
doc.save(output_path)
print(f"✅ Word剧本已保存到：{output_path}")
